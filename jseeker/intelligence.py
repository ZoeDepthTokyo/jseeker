"""JD corpus intelligence — aggregate patterns, ideal candidate profile, salary insight."""

from __future__ import annotations

import json
import logging
from collections import Counter
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def aggregate_jd_corpus(db_path=None) -> dict:
    """Aggregate all parsed JDs from jd_cache + salary from applications.

    Returns dict with frequency-ranked signals across the corpus:
        top_keywords, top_skills, top_requirements, salary_percentiles,
        experience_distribution, remote_policy_breakdown, culture_signals,
        total_jds, markets

    Args:
        db_path: Optional path override for the SQLite database.

    Returns:
        Dict with aggregated corpus statistics.
    """
    import sqlite3
    from jseeker.tracker import tracker_db

    db = db_path or tracker_db.db_path
    conn = sqlite3.connect(str(db))
    conn.row_factory = sqlite3.Row

    # --- Parse all cached JDs ---
    rows = conn.execute("SELECT parsed_json FROM jd_cache WHERE parsed_json IS NOT NULL").fetchall()

    keyword_counter: Counter = Counter()
    skill_counter: Counter = Counter()
    requirement_counter: Counter = Counter()
    culture_counter: Counter = Counter()
    markets: list[str] = []
    exp_list: list[str] = []
    remote_list: list[str] = []

    for row in rows:
        try:
            data = json.loads(row["parsed_json"])
        except (json.JSONDecodeError, TypeError):
            continue

        for kw in data.get("ats_keywords", []):
            keyword_counter[kw.lower()] += 1

        for cs in data.get("culture_signals", []):
            culture_counter[cs.lower()] += 1

        for req in data.get("requirements", []):
            text = req.get("text", "") if isinstance(req, dict) else str(req)
            for kw in req.get("keywords", []) if isinstance(req, dict) else []:
                skill_counter[kw.lower()] += 1
            if text:
                requirement_counter[text[:80]] += 1

        if data.get("market"):
            markets.append(data["market"])
        if data.get("role_exp"):
            exp_list.append(data["role_exp"])
        if data.get("remote_policy"):
            remote_list.append(data["remote_policy"])

    # --- Salary from applications ---
    salary_rows = conn.execute(
        "SELECT salary_min, salary_max, market FROM applications "
        "WHERE salary_min IS NOT NULL AND salary_max IS NOT NULL"
    ).fetchall()
    conn.close()

    salary_values = [
        (r["salary_min"] + r["salary_max"]) // 2
        for r in salary_rows
        if r["salary_min"] and r["salary_max"]
    ]
    salary_by_market: dict[str, list[int]] = {}
    for r in salary_rows:
        if r["salary_min"] and r["salary_max"]:
            mid = (r["salary_min"] + r["salary_max"]) // 2
            salary_by_market.setdefault(r["market"] or "unknown", []).append(mid)

    def _percentiles(vals: list[int]) -> dict:
        if not vals:
            return {}
        s = sorted(vals)
        n = len(s)
        return {
            "p25": s[n // 4],
            "p50": s[n // 2],
            "p75": s[3 * n // 4],
            "count": n,
        }

    return {
        "total_jds": len(rows),
        "top_keywords": keyword_counter.most_common(25),
        "top_skills": skill_counter.most_common(20),
        "top_requirements": requirement_counter.most_common(10),
        "culture_signals": culture_counter.most_common(15),
        "salary_percentiles": _percentiles(salary_values),
        "salary_by_market": {m: _percentiles(v) for m, v in salary_by_market.items()},
        "experience_distribution": dict(Counter(exp_list).most_common(10)),
        "remote_policy_breakdown": dict(Counter(remote_list).most_common(5)),
        "markets": list(set(markets)),
    }


def generate_ideal_candidate_brief(
    parsed_jd,
    adapted_resume,
    aggregate: dict,
    db_path=None,
) -> "IntelligenceReport":
    """Generate an ideal candidate brief using Sonnet. Caches result by JD hash.

    Cost: ~$0.015 Sonnet. Cache hit: $0.

    Args:
        parsed_jd: ParsedJD instance from the pipeline.
        adapted_resume: AdaptedResume instance (can be empty).
        aggregate: Output of aggregate_jd_corpus().
        db_path: Optional database path override.

    Returns:
        IntelligenceReport with ideal_profile, strengths, gaps, and salary data.
    """
    from jseeker.models import IntelligenceReport
    from jseeker.tracker import tracker_db
    from jseeker.llm import llm

    jd_hash = getattr(parsed_jd, "_pruned_hash", None) or (parsed_jd.title + parsed_jd.company)[:64]

    # Cache check
    cached = tracker_db.get_intelligence(jd_hash)
    if cached:
        return IntelligenceReport(
            jd_hash=jd_hash,
            ideal_profile=cached.get("ideal_profile", ""),
            strengths=json.loads(cached.get("strengths", "[]")),
            gaps=json.loads(cached.get("gaps", "[]")),
            salary_angle=cached.get("salary_angle", ""),
            keyword_coverage=cached.get("keyword_coverage", 0.0),
        )

    # Build prompt
    prompt_path = Path(__file__).parent.parent / "data" / "prompts" / "ideal_candidate.txt"
    template = prompt_path.read_text(encoding="utf-8")

    requirements = ""
    if hasattr(parsed_jd, "requirements") and parsed_jd.requirements:
        requirements = "\n".join(f"- {r.text}" for r in parsed_jd.requirements[:8])
    ats_keywords = ", ".join((parsed_jd.ats_keywords or [])[:15])
    culture_signals = ", ".join((parsed_jd.culture_signals or [])[:8])
    adapted_summary = getattr(adapted_resume, "summary", "") or ""
    skills_covered = (
        ", ".join(getattr(adapted_resume, "skills", [])[:10])
        if hasattr(adapted_resume, "skills")
        else ""
    )

    jd_kws = {kw for kw in (parsed_jd.ats_keywords or [])}
    resume_kws = set((getattr(adapted_resume, "skills", None) or []))
    missing = list(jd_kws - resume_kws)[:10]
    keyword_coverage = len(jd_kws & resume_kws) / max(len(jd_kws), 1)

    top_agg_kws = ", ".join(kw for kw, _ in aggregate.get("top_keywords", [])[:15])

    prompt = template.format(
        job_title=parsed_jd.title or "this role",
        company=parsed_jd.company or "this company",
        requirements=requirements or "See job description",
        ats_keywords=ats_keywords or "Not specified",
        culture_signals=culture_signals or "Professional",
        aggregate_top_keywords=top_agg_kws or "Not available",
        adapted_summary=adapted_summary[:500],
        skills_covered=skills_covered or "Various",
        missing_keywords=", ".join(missing) or "None identified",
    )

    raw = llm.call_sonnet(prompt, task="jd_intelligence")

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        import re

        m = re.search(r"\{.*\}", raw, re.DOTALL)
        data = json.loads(m.group()) if m else {}

    report = IntelligenceReport(
        jd_hash=jd_hash,
        ideal_profile=data.get("ideal_profile", ""),
        strengths=data.get("strengths", []),
        gaps=data.get("gaps", []),
        salary_angle=data.get("salary_angle", ""),
        keyword_coverage=keyword_coverage,
        salary_insight=_build_salary_insight(parsed_jd, aggregate),
    )

    tracker_db.save_intelligence(
        jd_hash,
        {
            "ideal_profile": report.ideal_profile,
            "strengths": report.strengths,
            "gaps": report.gaps,
            "salary_angle": report.salary_angle,
            "keyword_coverage": report.keyword_coverage,
        },
    )

    return report


def _build_salary_insight(parsed_jd, aggregate: dict) -> dict:
    """Build salary insight from SQL-derived corpus data. No LLM call.

    Args:
        parsed_jd: ParsedJD instance with market/salary fields.
        aggregate: Output of aggregate_jd_corpus().

    Returns:
        Dict with salary insight data or unavailability message.
    """
    market = getattr(parsed_jd, "market", "us") or "us"
    by_market = aggregate.get("salary_by_market", {})
    market_data = by_market.get(market, {})
    global_data = aggregate.get("salary_percentiles", {})

    jd_max = getattr(parsed_jd, "salary_max", None)
    target = jd_max or market_data.get("p75") or global_data.get("p75")

    if not target:
        return {
            "available": False,
            "message": "Insufficient salary data. Add more applications with salary info.",
        }

    return {
        "available": True,
        "target_ask": target,
        "market": market,
        "market_p50": market_data.get("p50"),
        "market_p75": market_data.get("p75"),
        "data_points": market_data.get("count", 0),
        "recommendation": (
            f"Ask for ${target:,} based on {market_data.get('count', 0)} "
            f"data points in {market.upper()} market."
        ),
    }


def export_profile_docx(report: "IntelligenceReport", output_path: Path) -> Path:
    """Export IntelligenceReport to a formatted DOCX file.

    Args:
        report: IntelligenceReport with profile data.
        output_path: Destination path for the DOCX file.

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Ideal Candidate Blueprint", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ideal Profile
    doc.add_heading("Ideal Candidate Profile", level=1)
    doc.add_paragraph(report.ideal_profile)

    # Strengths
    if report.strengths:
        doc.add_heading("Your Strengths for This Role", level=1)
        for s in report.strengths:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(s)
            run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)  # forest green

    # Gaps
    if report.gaps:
        doc.add_heading("Gaps to Address", level=1)
        for g in report.gaps:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(g)
            run.font.color.rgb = RGBColor(0xD2, 0x69, 0x1E)  # dark orange

    # Salary
    if report.salary_angle:
        doc.add_heading("Salary Negotiation Angle", level=1)
        doc.add_paragraph(report.salary_angle)

    # Coverage metric
    doc.add_heading("Keyword Coverage", level=2)
    doc.add_paragraph(f"{report.keyword_coverage:.0%} of JD keywords covered by your resume.")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
