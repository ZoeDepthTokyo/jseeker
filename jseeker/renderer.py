"""jSeeker Renderer — HTML→PDF (Playwright subprocess) + DOCX (python-docx)."""

from __future__ import annotations

import subprocess
import sys
import tempfile
import yaml
from datetime import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

from jseeker.models import AdaptedResume

# Section labels for bilingual support
SECTION_LABELS = {
    "en": {
        "contact": "CONTACT",
        "online": "ONLINE",
        "skills": "SKILLS",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "awards": "AWARDS",
        "languages": "LANGUAGES",
        "summary": "SUMMARY",
        "experience": "EXPERIENCE",
        "early_career": "EARLY CAREER",
    },
    "es": {
        "contact": "CONTACTO",
        "online": "EN LINEA",
        "skills": "HABILIDADES",
        "education": "EDUCACION",
        "certifications": "CERTIFICACIONES",
        "awards": "PREMIOS",
        "languages": "IDIOMAS",
        "summary": "RESUMEN PROFESIONAL",
        "experience": "EXPERIENCIA PROFESIONAL",
        "early_career": "CARRERA TEMPRANA",
    },
}


def _get_templates_env() -> Environment:
    """Create Jinja2 environment for templates."""
    from config import settings
    return Environment(
        loader=FileSystemLoader(str(settings.templates_dir)),
        autoescape=True,
    )


def _render_html(adapted: AdaptedResume, template_name: str, language: str = "en") -> str:
    """Render resume data into HTML using a Jinja2 template.

    Args:
        adapted: Adapted resume content.
        template_name: Name of the Jinja2 template to use.
        language: Language code for section labels ("en" or "es").

    Returns:
        Rendered HTML string.
    """
    env = _get_templates_env()
    template = env.get_template(template_name)

    # Format dates for display
    experiences = []
    for exp in adapted.experience_blocks:
        end = exp.get("end")
        if end is None:
            end_display = "Present" if language == "en" else "Presente"
        else:
            end_display = _format_date(end)
        experiences.append({
            **exp,
            "start_display": _format_date(exp.get("start", "")),
            "end_display": end_display,
        })

    education = []
    for edu in adapted.education:
        education.append({
            "institution": edu.institution,
            "degree": edu.degree or "",
            "field": edu.field,
            "start": edu.start,
            "end": edu.end,
        })

    # Get section labels for the specified language
    section_labels = SECTION_LABELS.get(language, SECTION_LABELS["en"])

    return template.render(
        name=adapted.contact.full_name,
        target_title=adapted.target_title,
        email=adapted.contact.email,
        phone=adapted.contact.phone,
        website=adapted.contact.website,
        linkedin=adapted.contact.linkedin,
        locations=adapted.contact.locations,
        languages=adapted.contact.languages,
        summary=adapted.summary,
        experiences=experiences,
        skills=adapted.skills_ordered,
        education=education,
        certifications=adapted.certifications,
        awards=adapted.awards,
        early_career=adapted.early_career,
        section_labels=section_labels,
    )


def _format_date(date_str: str) -> str:
    """Convert YYYY-MM or YYYY to MM/YYYY format."""
    if not date_str:
        return ""
    parts = date_str.split("-")
    if len(parts) == 2:
        return f"{parts[1]}/{parts[0]}"
    return date_str


def _html_to_pdf_sync(html: str, output_path: Path) -> Path:
    """Convert HTML to PDF using Playwright in a subprocess.

    Playwright's sync/async APIs both conflict with Streamlit's event loop
    on Windows, so we spawn a separate Python process to do the rendering.

    DEPRECATED: Use browser_manager.html_to_pdf_fast() for 90% faster rendering.
    This function kept for backwards compatibility.
    """
    # Write HTML to a temp file
    html_tmp = Path(tempfile.mktemp(suffix=".html"))
    html_tmp.write_text(html, encoding="utf-8")

    pdf_str = str(output_path).replace("\\", "\\\\")
    html_uri = f"file:///{html_tmp.as_posix()}"

    script = f"""
from playwright.sync_api import sync_playwright
with sync_playwright() as p:
    browser = p.chromium.launch()
    page = browser.new_page()
    page.goto("{html_uri}")
    page.pdf(
        path=r"{pdf_str}",
        format="Letter",
        print_background=True,
        margin={{"top": "0.5in", "bottom": "0.5in", "left": "0.5in", "right": "0.5in"}},
    )
    browser.close()
"""

    try:
        result = subprocess.run(
            [sys.executable, "-c", script],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"PDF generation failed: {result.stderr[:500]}")
    finally:
        html_tmp.unlink(missing_ok=True)

    return output_path


def _html_to_pdf_fast(html: str, output_path: Path) -> Path:
    """Convert HTML to PDF using persistent browser (90% faster after first call).

    First call: 5-15s (browser startup)
    Subsequent calls: 1-2s

    Falls back to _html_to_pdf_sync if persistent browser fails.
    """
    try:
        from jseeker.browser_manager import html_to_pdf_fast
        return html_to_pdf_fast(html, output_path)
    except Exception as e:
        # Fallback to slow method if persistent browser fails
        import warnings
        warnings.warn(f"Fast PDF rendering failed ({e}), falling back to slow method")
        return _html_to_pdf_sync(html, output_path)


def render_pdf(
    adapted: AdaptedResume,
    output_path: Path,
    two_column: bool = True,
    language: str = "en",
    use_fast_renderer: bool = True,
) -> Path:
    """Render resume to PDF via HTML + Playwright.

    Args:
        adapted: Adapted resume content.
        output_path: Where to save the PDF.
        two_column: Use two-column (visual) or single-column (ATS-safe) template.
        language: Language code for section labels ("en" or "es").
        use_fast_renderer: Use persistent browser (90% faster after first call).

    Returns:
        Path to generated PDF.
    """
    template_name = "two_column.html" if two_column else "single_column.html"
    html = _render_html(adapted, template_name, language=language)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if use_fast_renderer:
        return _html_to_pdf_fast(html, output_path)
    else:
        return _html_to_pdf_sync(html, output_path)


def render_docx(adapted: AdaptedResume, output_path: Path, language: str = "en") -> Path:
    """Render resume to DOCX (single-column ATS-safe).

    Uses python-docx for maximum ATS compatibility.
    Compact 2-page layout with blue section headers.

    Args:
        adapted: Adapted resume content.
        output_path: Path to save the DOCX file.
        language: Language code for section labels ("en" or "es").

    Returns:
        Path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Get section labels for the specified language
    section_labels = SECTION_LABELS.get(language, SECTION_LABELS["en"])

    doc = Document()

    # Set default font to 10pt for compact layout
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Header: Name (blue, bold, 16pt)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(adapted.contact.full_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0x2B, 0x57, 0x97)

    # Header: Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run(adapted.target_title)
    title_run.font.size = Pt(12)

    # Contact line
    contact_parts = []
    if adapted.contact.email:
        contact_parts.append(adapted.contact.email)
    if adapted.contact.phone:
        contact_parts.append(adapted.contact.phone)
    if adapted.contact.locations:
        contact_parts.append(adapted.contact.locations[0])
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(2)
    contact_para.add_run(" | ".join(contact_parts)).font.size = Pt(10)

    # Links line
    links = []
    if adapted.contact.website:
        links.append(adapted.contact.website)
    if adapted.contact.linkedin:
        links.append(adapted.contact.linkedin)
    if links:
        links_para = doc.add_paragraph()
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_para.paragraph_format.space_after = Pt(4)
        links_para.add_run(" | ".join(links)).font.size = Pt(10)

    # Summary
    _add_section_header(doc, section_labels["summary"])
    summary_para = doc.add_paragraph(adapted.summary)
    summary_para.paragraph_format.space_after = Pt(4)

    # Experience
    _add_section_header(doc, section_labels["experience"])
    for exp in adapted.experience_blocks:
        end = exp.get("end")
        end_display = ("Present" if language == "en" else "Presente") if end is None else _format_date(end)
        start_display = _format_date(exp.get("start", ""))

        # Check if this is a condensed entry
        is_condensed = exp.get("condensed", False)

        role_para = doc.add_paragraph()
        role_para.paragraph_format.space_before = Pt(0)
        role_para.paragraph_format.space_after = Pt(0)
        role_run = role_para.add_run(f"{exp['role']}")
        role_run.bold = True
        role_para.add_run(f" — {exp['company']}")
        if is_condensed:
            role_run.font.size = Pt(9.5)
            role_para.runs[-1].font.size = Pt(9.5)

        date_para = doc.add_paragraph()
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(1)
        date_run = date_para.add_run(f"{start_display} – {end_display}")
        if exp.get("location"):
            date_para.add_run(f" | {exp['location']}")
        if is_condensed:
            for run in date_para.runs:
                run.font.size = Pt(9.5)

        # Limit bullets for condensed entries
        bullets = exp.get("bullets", [])
        if is_condensed and len(bullets) > 3:
            bullets = bullets[:3]

        for bullet in bullets:
            bullet_para = doc.add_paragraph(style="List Bullet")
            bullet_para.paragraph_format.space_before = Pt(0)
            bullet_para.paragraph_format.space_after = Pt(1)
            bullet_run = bullet_para.add_run(bullet)
            if is_condensed:
                bullet_run.font.size = Pt(9.5)
            else:
                bullet_run.font.size = Pt(10)

    # Skills
    _add_section_header(doc, section_labels["skills"])
    for skill_group in adapted.skills_ordered:
        skills_str = ", ".join(skill_group.get("skills", []))
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        cat_run = para.add_run(f"{skill_group['category']}: ")
        cat_run.bold = True
        para.add_run(skills_str)

    # Education
    _add_section_header(doc, section_labels["education"])
    for edu in adapted.education:
        degree = edu.degree or ""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2)
        edu_run = para.add_run(f"{degree} {edu.field}")
        edu_run.bold = True
        para.add_run(f" — {edu.institution} ({edu.start}–{edu.end})")

    # Awards
    if adapted.awards:
        _add_section_header(doc, section_labels["awards"])
        for award in adapted.awards:
            award_para = doc.add_paragraph(f"{award.name}", style="List Bullet")
            award_para.paragraph_format.space_before = Pt(0)
            award_para.paragraph_format.space_after = Pt(1)

    # Certifications
    if adapted.certifications:
        _add_section_header(doc, section_labels["certifications"])
        cert_names = [c.name for c in adapted.certifications]
        cert_para = doc.add_paragraph(", ".join(cert_names))
        cert_para.paragraph_format.space_before = Pt(0)
        cert_para.paragraph_format.space_after = Pt(4)

    # Early Career (compact one-liners)
    if adapted.early_career:
        _add_section_header(doc, section_labels["early_career"])
        for entry in adapted.early_career:
            role = entry.get("role", "")
            company = entry.get("company", "")
            para = doc.add_paragraph(f"{role} — {company}", style="List Bullet")
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)
            for run in para.runs:
                run.font.size = Pt(9.5)

    # Languages
    if hasattr(adapted.contact, 'languages') and adapted.contact.languages:
        _add_section_header(doc, section_labels["languages"])
        lang_para = doc.add_paragraph(", ".join(adapted.contact.languages))
        lang_para.paragraph_format.space_before = Pt(0)
        lang_para.paragraph_format.space_after = Pt(4)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _add_bottom_border(paragraph, color: str = "2B5797") -> None:
    """Add a thin bottom border to a paragraph.

    Args:
        paragraph: The paragraph to add the border to.
        color: Hex color string (without #) for the border.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 0.5pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)


def _add_section_header(doc, text: str) -> None:
    """Add a blue section header with bottom border.

    Args:
        doc: Document object.
        text: Section header text.
    """
    from docx.shared import Pt, RGBColor
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x97)  # Blue color

    # Add spacing and bottom border
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    _add_bottom_border(para, "2B5797")


def _get_display_name() -> str:
    """Read display_name from contact.yaml."""
    from config import settings
    contact_path = settings.resume_blocks_dir / "contact.yaml"
    data = yaml.safe_load(contact_path.read_text(encoding="utf-8"))
    return data.get("contact", {}).get("display_name", "Resume")


def _sanitize(text: str, max_len: int = 30) -> str:
    """Sanitize text for use in filenames."""
    clean = "".join(c for c in text if c.isalnum() or c in " -_")
    return clean.strip().replace(" ", "_")[:max_len]


def _get_next_version(folder: Path, base_name: str) -> int:
    """Find next available version number in a folder."""
    if not folder.exists():
        return 1
    existing = list(folder.glob(f"{base_name}_v*.pdf")) + list(folder.glob(f"{base_name}_v*.docx"))
    if not existing:
        return 1
    versions = []
    for f in existing:
        stem = f.stem
        try:
            v = int(stem.rsplit("_v", 1)[1])
            versions.append(v)
        except (ValueError, IndexError):
            pass
    return max(versions, default=0) + 1


def generate_output(
    adapted: AdaptedResume,
    company: str,
    role: str,
    output_dir: Path = None,
    formats: list[str] = None,
    filename_override: str = None,
    language: str = "en",
) -> dict[str, Path]:
    """Generate all output files for a resume.

    Naming: {display_name}_{role}_{company}_v{N}.{ext}
    Folder: output/{company}/

    Args:
        adapted: Adapted resume content.
        company: Company name (for folder naming).
        role: Role title (for folder naming).
        output_dir: Base output directory.
        formats: List of formats to generate ("pdf", "docx", or both).
        filename_override: Override the base filename (sanitized, version still appended).
        language: Language code for section labels ("en" or "es").

    Returns:
        Dict of {format: path} for generated files.
    """
    if output_dir is None:
        from config import settings
        output_dir = settings.output_dir

    if formats is None:
        formats = ["pdf", "docx"]

    # Build folder and filename
    display_name = _get_display_name()
    safe_name = _sanitize(display_name)
    safe_company = _sanitize(company)
    safe_role = _sanitize(role)

    folder = output_dir / safe_company
    folder.mkdir(parents=True, exist_ok=True)

    base_name = f"{safe_name}_{safe_role}_{safe_company}"

    if filename_override:
        base_name = _sanitize(filename_override, max_len=80)

    version = _get_next_version(folder, base_name)

    results = {}

    if "pdf" in formats:
        pdf_path = folder / f"{base_name}_v{version}.pdf"
        render_pdf(adapted, pdf_path, two_column=True, language=language)
        results["pdf"] = pdf_path

    if "docx" in formats:
        docx_path = folder / f"{base_name}_v{version}.docx"
        render_docx(adapted, docx_path, language=language)
        results["docx"] = docx_path

    return results
