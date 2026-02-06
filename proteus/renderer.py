"""PROTEUS Renderer — HTML→PDF (Playwright) + DOCX (python-docx)."""

from __future__ import annotations

import asyncio
from datetime import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

from proteus.models import AdaptedResume


def _get_templates_env() -> Environment:
    """Create Jinja2 environment for templates."""
    from config import settings
    return Environment(
        loader=FileSystemLoader(str(settings.templates_dir)),
        autoescape=True,
    )


def _render_html(adapted: AdaptedResume, template_name: str) -> str:
    """Render resume data into HTML using a Jinja2 template."""
    env = _get_templates_env()
    template = env.get_template(template_name)

    # Format dates for display
    experiences = []
    for exp in adapted.experience_blocks:
        end = exp.get("end")
        if end is None:
            end_display = "Present"
        else:
            end_display = _format_date(end)
        experiences.append({
            **exp,
            "start_display": _format_date(exp.get("start", "")),
            "end_display": end_display,
        })

    education = []
    for edu in adapted.education:
        education.append({
            "institution": edu.institution,
            "degree": edu.degree or "",
            "field": edu.field,
            "start": edu.start,
            "end": edu.end,
        })

    return template.render(
        name=adapted.contact.full_name,
        target_title=adapted.target_title,
        email=adapted.contact.email,
        phone=adapted.contact.phone,
        website=adapted.contact.website,
        linkedin=adapted.contact.linkedin,
        locations=adapted.contact.locations,
        languages=adapted.contact.languages,
        summary=adapted.summary,
        experiences=experiences,
        skills=adapted.skills_ordered,
        education=education,
        certifications=adapted.certifications,
        awards=adapted.awards,
        early_career=adapted.early_career,
    )


def _format_date(date_str: str) -> str:
    """Convert YYYY-MM or YYYY to MM/YYYY format."""
    if not date_str:
        return ""
    parts = date_str.split("-")
    if len(parts) == 2:
        return f"{parts[1]}/{parts[0]}"
    return date_str


async def _html_to_pdf_async(html: str, output_path: Path) -> Path:
    """Convert HTML to PDF using Playwright Chromium."""
    from playwright.async_api import async_playwright

    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.set_content(html, wait_until="networkidle")
        await page.pdf(
            path=str(output_path),
            format="Letter",
            margin={"top": "0.5in", "bottom": "0.5in", "left": "0.5in", "right": "0.5in"},
            print_background=True,
        )
        await browser.close()

    return output_path


def render_pdf(
    adapted: AdaptedResume,
    output_path: Path,
    two_column: bool = True,
) -> Path:
    """Render resume to PDF via HTML + Playwright.

    Args:
        adapted: Adapted resume content.
        output_path: Where to save the PDF.
        two_column: Use two-column (visual) or single-column (ATS-safe) template.

    Returns:
        Path to generated PDF.
    """
    template_name = "two_column.html" if two_column else "single_column.html"
    html = _render_html(adapted, template_name)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    return asyncio.run(_html_to_pdf_async(html, output_path))


def render_docx(adapted: AdaptedResume, output_path: Path) -> Path:
    """Render resume to DOCX (single-column ATS-safe).

    Uses python-docx for maximum ATS compatibility.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header: Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(adapted.contact.full_name)
    name_run.bold = True
    name_run.font.size = Pt(16)

    # Header: Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(adapted.target_title)
    title_run.font.size = Pt(12)

    # Contact line
    contact_parts = []
    if adapted.contact.email:
        contact_parts.append(adapted.contact.email)
    if adapted.contact.phone:
        contact_parts.append(adapted.contact.phone)
    if adapted.contact.locations:
        contact_parts.append(adapted.contact.locations[0])
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run(" | ".join(contact_parts)).font.size = Pt(10)

    # Links line
    links = []
    if adapted.contact.website:
        links.append(adapted.contact.website)
    if adapted.contact.linkedin:
        links.append(adapted.contact.linkedin)
    if links:
        links_para = doc.add_paragraph()
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_para.add_run(" | ".join(links)).font.size = Pt(10)

    # Summary
    _add_section_header(doc, "SUMMARY")
    doc.add_paragraph(adapted.summary)

    # Experience
    _add_section_header(doc, "EXPERIENCE")
    for exp in adapted.experience_blocks:
        end = exp.get("end")
        end_display = "Present" if end is None else _format_date(end)
        start_display = _format_date(exp.get("start", ""))

        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"{exp['role']}")
        role_run.bold = True
        role_para.add_run(f" — {exp['company']}")

        date_para = doc.add_paragraph()
        date_para.add_run(f"{start_display} – {end_display}")
        if exp.get("location"):
            date_para.add_run(f" | {exp['location']}")

        for bullet in exp.get("bullets", []):
            bullet_para = doc.add_paragraph(style="List Bullet")
            bullet_para.add_run(bullet)

    # Skills
    _add_section_header(doc, "SKILLS")
    for skill_group in adapted.skills_ordered:
        skills_str = ", ".join(skill_group.get("skills", []))
        para = doc.add_paragraph()
        cat_run = para.add_run(f"{skill_group['category']}: ")
        cat_run.bold = True
        para.add_run(skills_str)

    # Education
    _add_section_header(doc, "EDUCATION")
    for edu in adapted.education:
        degree = edu.degree or ""
        para = doc.add_paragraph()
        edu_run = para.add_run(f"{degree} {edu.field}")
        edu_run.bold = True
        para.add_run(f" — {edu.institution} ({edu.start}–{edu.end})")

    # Awards
    if adapted.awards:
        _add_section_header(doc, "AWARDS")
        for award in adapted.awards:
            doc.add_paragraph(f"{award.name}", style="List Bullet")

    # Certifications
    if adapted.certifications:
        _add_section_header(doc, "CERTIFICATIONS")
        cert_names = [c.name for c in adapted.certifications]
        doc.add_paragraph(", ".join(cert_names))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _add_section_header(doc, text: str) -> None:
    """Add a section header with underline."""
    from docx.shared import Pt
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    # Add a thin line below (paragraph border)
    para.paragraph_format.space_after = Pt(4)


def generate_output(
    adapted: AdaptedResume,
    company: str,
    role: str,
    output_dir: Path = None,
    formats: list[str] = None,
) -> dict[str, Path]:
    """Generate all output files for a resume.

    Args:
        adapted: Adapted resume content.
        company: Company name (for folder naming).
        role: Role title (for folder naming).
        output_dir: Base output directory.
        formats: List of formats to generate ("pdf", "docx", or both).

    Returns:
        Dict of {format: path} for generated files.
    """
    if output_dir is None:
        from config import settings
        output_dir = settings.output_dir

    if formats is None:
        formats = ["pdf", "docx"]

    # Create output folder
    date_str = datetime.now().strftime("%Y%m%d")
    safe_company = "".join(c for c in company if c.isalnum() or c in " -_")[:30].strip()
    safe_role = "".join(c for c in role if c.isalnum() or c in " -_")[:30].strip()
    folder = output_dir / f"{date_str}_{safe_company}_{safe_role}"
    folder.mkdir(parents=True, exist_ok=True)

    results = {}

    if "pdf" in formats:
        pdf_path = folder / "resume_v1.pdf"
        render_pdf(adapted, pdf_path, two_column=True)
        results["pdf"] = pdf_path

    if "docx" in formats:
        docx_path = folder / "resume_v1.docx"
        render_docx(adapted, docx_path)
        results["docx"] = docx_path

    return results
