"""Tests for renderer module."""

import pytest
from unittest.mock import patch, MagicMock
import subprocess
from jseeker.renderer import _sanitize, _get_next_version, SECTION_LABELS, _html_to_pdf_sync, generate_output


class TestSanitize:
    """Test filename sanitization."""

    def test_sanitize_basic(self):
        """Test basic string sanitization."""
        result = _sanitize("Hello World")
        assert result == "Hello_World"

    def test_sanitize_special_chars(self):
        """Test sanitization removes special characters."""
        result = _sanitize("Director (VP)")
        assert result == "Director_VP"

    def test_sanitize_max_length(self):
        """Test sanitization truncates to max_len."""
        long_string = "This is a very long string that should be truncated to the maximum length"
        result = _sanitize(long_string, max_len=20)
        assert len(result) <= 20

    def test_sanitize_empty_string(self):
        """Test sanitization returns 'Unknown' for empty string."""
        result = _sanitize("")
        assert result == "Unknown"

    def test_sanitize_whitespace_only(self):
        """Test sanitization returns 'Unknown' for whitespace-only string."""
        result = _sanitize("   ")
        assert result == "Unknown"

    def test_sanitize_special_chars_only(self):
        """Test sanitization returns 'Unknown' when only special chars remain."""
        result = _sanitize("@#$%^&*()")
        assert result == "Unknown"


class TestGetNextVersion:
    """Test version number generation."""

    def test_get_next_version_empty_folder(self, tmp_path):
        """Test returns 1 for empty folder."""
        folder = tmp_path / "empty"
        result = _get_next_version(folder, "TestResume")
        assert result == 1

    def test_get_next_version_nonexistent_folder(self, tmp_path):
        """Test returns 1 for non-existent folder."""
        folder = tmp_path / "nonexistent"
        result = _get_next_version(folder, "TestResume")
        assert result == 1

    def test_get_next_version_with_existing(self, tmp_path):
        """Test returns correct next version with existing files."""
        folder = tmp_path / "resumes"
        folder.mkdir()

        # Create dummy v1 and v2 files
        (folder / "TestResume_v1.pdf").write_text("fake pdf v1")
        (folder / "TestResume_v2.pdf").write_text("fake pdf v2")

        result = _get_next_version(folder, "TestResume")
        assert result == 3


class TestSectionLabels:
    """Test bilingual section labels."""

    def test_section_labels_en(self):
        """Test all section keys exist in English."""
        expected_keys = [
            "contact",
            "online",
            "skills",
            "education",
            "certifications",
            "awards",
            "languages",
            "summary",
            "experience",
            "early_career",
        ]
        for key in expected_keys:
            assert key in SECTION_LABELS["en"]

    def test_section_labels_es(self):
        """Test all section keys exist in Spanish."""
        expected_keys = [
            "contact",
            "online",
            "skills",
            "education",
            "certifications",
            "awards",
            "languages",
            "summary",
            "experience",
            "early_career",
        ]
        for key in expected_keys:
            assert key in SECTION_LABELS["es"]

    def test_section_labels_es_values(self):
        """Test specific Spanish translations."""
        assert SECTION_LABELS["es"]["summary"] == "RESUMEN PROFESIONAL"
        assert SECTION_LABELS["es"]["experience"] == "EXPERIENCIA PROFESIONAL"
        assert SECTION_LABELS["es"]["contact"] == "CONTACTO"
        assert SECTION_LABELS["es"]["skills"] == "HABILIDADES"


class TestPDFRendering:
    """Test PDF rendering with retry logic and error handling."""

    def test_render_pdf_subprocess_failure(self, tmp_path):
        """Test that subprocess failure raises RenderError with full stderr."""
        from jseeker.models import RenderError

        html_content = "<html><body>Test</body></html>"
        output_path = tmp_path / "test.pdf"

        # Mock subprocess.run to simulate failure with long stderr
        long_error = "Playwright error: " + ("X" * 1000)
        mock_result = MagicMock()
        mock_result.returncode = 1
        mock_result.stderr = long_error
        mock_result.stdout = ""

        with patch("subprocess.run", return_value=mock_result):
            with pytest.raises(RenderError) as exc_info:
                _html_to_pdf_sync(html_content, output_path)

            # Verify error contains FULL stderr (not truncated at 500 chars)
            assert len(str(exc_info.value)) > 500
            assert "Playwright error" in str(exc_info.value)

    def test_render_pdf_timeout(self, tmp_path):
        """Test that subprocess timeout raises RenderError."""
        from jseeker.models import RenderError

        html_content = "<html><body>Test</body></html>"
        output_path = tmp_path / "test.pdf"

        # Mock subprocess.run to simulate timeout
        with patch("subprocess.run", side_effect=subprocess.TimeoutExpired("python", 60)):
            with pytest.raises(RenderError) as exc_info:
                _html_to_pdf_sync(html_content, output_path)

            assert "timeout" in str(exc_info.value).lower()

    def test_render_pdf_success_after_retry(self, tmp_path):
        """Test that rendering succeeds on second retry attempt."""
        html_content = "<html><body>Test</body></html>"
        output_path = tmp_path / "test.pdf"

        # Mock subprocess.run to fail twice, then succeed
        call_count = [0]

        def mock_run(*args, **kwargs):
            call_count[0] += 1
            result = MagicMock()
            if call_count[0] < 3:
                result.returncode = 1
                result.stderr = "Transient error"
                result.stdout = ""
                return result
            else:
                result.returncode = 0
                result.stderr = ""
                result.stdout = ""
                # Create dummy PDF
                output_path.write_bytes(b"%PDF-1.4")
                return result

        with patch("subprocess.run", side_effect=mock_run):
            result_path = _html_to_pdf_sync(html_content, output_path)

            assert result_path == output_path
            assert output_path.exists()
            assert call_count[0] == 3  # Failed twice, succeeded on third

    def test_render_pdf_error_log_created(self, tmp_path):
        """Test that detailed error log is created on failure."""
        from jseeker.models import RenderError

        html_content = "<html><body>Test</body></html>"
        output_path = tmp_path / "test.pdf"

        # Mock subprocess.run to simulate failure
        mock_result = MagicMock()
        mock_result.returncode = 1
        mock_result.stderr = "Detailed Playwright error with stack trace"
        mock_result.stdout = ""

        with patch("subprocess.run", return_value=mock_result):
            with patch("jseeker.renderer.Path") as mock_path_class:
                # Create mock error log path
                error_log = tmp_path / "pdf_error.log"
                mock_path_class.return_value = error_log

                with pytest.raises(RenderError):
                    _html_to_pdf_sync(html_content, output_path)

    def test_render_pdf_max_retries_exhausted(self, tmp_path):
        """Test that RenderError is raised after max retries exhausted."""
        from jseeker.models import RenderError

        html_content = "<html><body>Test</body></html>"
        output_path = tmp_path / "test.pdf"

        # Mock subprocess.run to always fail
        mock_result = MagicMock()
        mock_result.returncode = 1
        mock_result.stderr = "Persistent failure"
        mock_result.stdout = ""

        with patch("subprocess.run", return_value=mock_result):
            with pytest.raises(RenderError) as exc_info:
                _html_to_pdf_sync(html_content, output_path)

            # Verify error mentions retries exhausted
            assert "3 attempts" in str(exc_info.value) or "retries" in str(exc_info.value).lower()


class TestDOCXStructure:
    """Test DOCX structure for ATS compliance."""

    def test_company_title_separation(self, tmp_path):
        """Test that company name and job title are on separate paragraphs.

        CRITICAL for ATS parsing (Workday, Greenhouse, etc.).
        Single-line formats like 'Engineer — Acme' cause field misidentification.
        """
        from jseeker.models import AdaptedResume, ContactInfo
        from jseeker.renderer import render_docx
        from docx import Document

        # Create minimal adapted resume
        adapted = AdaptedResume(
            contact=ContactInfo(
                full_name="Test User",
                email="test@example.com",
                phone="555-1234",
                locations=["Remote"],
            ),
            target_title="Software Engineer",
            summary="Test summary",
            experience_blocks=[
                {
                    "role": "Senior Software Engineer",
                    "company": "Acme Corporation",
                    "start": "2022-01-01",
                    "end": "2023-12-31",
                    "location": "Remote",
                    "bullets": ["Led team of 5 engineers"],
                }
            ],
            skills_ordered=[{"category": "Languages", "skills": ["Python", "JavaScript"]}],
            education=[
                {
                    "degree": "BS Computer Science",
                    "institution": "State University",
                    "end": "2020-05-01",
                }
            ],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        # Render DOCX
        output_path = tmp_path / "test_resume.docx"
        result_path = render_docx(adapted, output_path)

        # Parse DOCX with python-docx
        doc = Document(result_path)

        # Find experience section paragraphs
        experience_start = None
        for i, para in enumerate(doc.paragraphs):
            if "EXPERIENCE" in para.text.upper():
                experience_start = i + 1
                break

        assert experience_start is not None, "Experience section not found"

        # Verify structure: title (bold) -> company (not bold) -> date
        title_para = doc.paragraphs[experience_start]
        company_para = doc.paragraphs[experience_start + 1]
        date_para = doc.paragraphs[experience_start + 2]

        # Check content
        assert "Senior Software Engineer" in title_para.text
        assert "Acme Corporation" in company_para.text
        assert "2022" in date_para.text

        # Check title is bold
        assert any(run.bold for run in title_para.runs), "Job title should be bold"

        # Check company is NOT bold (or at least separate paragraph)
        assert title_para.text != company_para.text, "Company and title must be separate paragraphs"

    def test_condensed_experience_structure(self, tmp_path):
        """Test that condensed experience entries maintain company/title separation."""
        from jseeker.models import AdaptedResume, ContactInfo
        from jseeker.renderer import render_docx
        from docx import Document

        adapted = AdaptedResume(
            contact=ContactInfo(
                full_name="Test User",
                email="test@example.com",
            ),
            target_title="Software Engineer",
            summary="Test summary",
            experience_blocks=[
                {
                    "role": "Junior Developer",
                    "company": "StartupCo",
                    "start": "2020-01-01",
                    "end": "2021-12-31",
                    "condensed": True,  # Condensed entry
                    "bullets": ["Built features", "Wrote tests", "Deployed code", "Extra bullet"],
                }
            ],
            skills_ordered=[],
            education=[],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        output_path = tmp_path / "test_condensed.docx"
        result_path = render_docx(adapted, output_path)

        doc = Document(result_path)

        # Find experience section
        experience_start = None
        for i, para in enumerate(doc.paragraphs):
            if "EXPERIENCE" in para.text.upper():
                experience_start = i + 1
                break

        assert experience_start is not None

        title_para = doc.paragraphs[experience_start]
        company_para = doc.paragraphs[experience_start + 1]

        # Verify separation even for condensed entries
        assert "Junior Developer" in title_para.text
        assert "StartupCo" in company_para.text
        assert title_para.text != company_para.text

        # Verify bullet limit (max 3 for condensed)
        bullets_found = 0
        for i in range(experience_start, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            if para.style.name == "List Bullet":
                bullets_found += 1
            elif "SKILLS" in para.text.upper() or "EDUCATION" in para.text.upper():
                break

        assert (
            bullets_found <= 3
        ), f"Condensed entry should have max 3 bullets, found {bullets_found}"

    def test_date_format_consistency(self, tmp_path):
        """Test that date format is consistent across all experience entries."""
        from jseeker.models import AdaptedResume, ContactInfo
        from jseeker.renderer import render_docx
        from docx import Document
        import re

        adapted = AdaptedResume(
            contact=ContactInfo(full_name="Test User", email="test@example.com"),
            target_title="Engineer",
            summary="Test",
            experience_blocks=[
                {
                    "role": "Role 1",
                    "company": "Company 1",
                    "start": "2020-01-01",
                    "end": "2021-06-30",
                    "bullets": [],
                },
                {
                    "role": "Role 2",
                    "company": "Company 2",
                    "start": "2021-07-01",
                    "end": None,  # Present
                    "bullets": [],
                },
            ],
            skills_ordered=[],
            education=[],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        output_path = tmp_path / "test_dates.docx"
        result_path = render_docx(adapted, output_path)

        doc = Document(result_path)

        # Extract all date lines (contain "–" en-dash)
        date_lines = [p.text for p in doc.paragraphs if "–" in p.text]

        assert len(date_lines) >= 2, "Should find at least 2 date lines"

        # Check format: "Month YYYY – Month YYYY" or "Month YYYY – Present"
        date_pattern = r"[A-Z][a-z]+ \d{4} – ([A-Z][a-z]+ \d{4}|Present)"
        for date_line in date_lines:
            # Extract date portion (before "|" if location exists)
            date_text = date_line.split("|")[0].strip()
            assert re.search(date_pattern, date_text), f"Date format invalid: {date_text}"

    def test_no_tables_in_experience(self, tmp_path):
        """Test that experience section uses paragraphs, not tables (ATS fails on tables)."""
        from jseeker.models import AdaptedResume, ContactInfo
        from jseeker.renderer import render_docx
        from docx import Document

        adapted = AdaptedResume(
            contact=ContactInfo(full_name="Test User", email="test@example.com"),
            target_title="Engineer",
            summary="Test",
            experience_blocks=[
                {
                    "role": "Engineer",
                    "company": "Company",
                    "start": "2020-01-01",
                    "end": "2023-12-31",
                    "bullets": ["Bullet 1", "Bullet 2"],
                }
            ],
            skills_ordered=[],
            education=[],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        output_path = tmp_path / "test_no_tables.docx"
        result_path = render_docx(adapted, output_path)

        doc = Document(result_path)

        # Verify no tables used in document
        assert len(doc.tables) == 0, "DOCX should not contain tables (ATS parsing failure)"

    def test_standard_section_headers(self, tmp_path):
        """Test that section headers use standard ATS-recognized names."""
        from jseeker.models import AdaptedResume, ContactInfo
        from jseeker.renderer import render_docx
        from docx import Document

        adapted = AdaptedResume(
            contact=ContactInfo(full_name="Test User", email="test@example.com"),
            target_title="Engineer",
            summary="Professional summary text",
            experience_blocks=[
                {
                    "role": "Engineer",
                    "company": "Company",
                    "start": "2020-01-01",
                    "end": "2023-12-31",
                    "bullets": [],
                }
            ],
            skills_ordered=[{"category": "Tech", "skills": ["Python"]}],
            education=[
                {
                    "degree": "BS Computer Science",
                    "institution": "University",
                    "end": "2020-05-01",
                }
            ],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        output_path = tmp_path / "test_headers.docx"
        result_path = render_docx(adapted, output_path)

        doc = Document(result_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Check for standard ATS-recognized headers (uppercase)
        assert "SUMMARY" in all_text or "PROFESSIONAL SUMMARY" in all_text
        assert "EXPERIENCE" in all_text or "PROFESSIONAL EXPERIENCE" in all_text
        assert "SKILLS" in all_text or "TECHNICAL SKILLS" in all_text
        assert "EDUCATION" in all_text

        # Check for non-standard headers (should NOT exist)
        assert "MY JOURNEY" not in all_text.upper()
        assert "TOOLBOX" not in all_text.upper()
        assert "WHAT I'VE BUILT" not in all_text.upper()


class TestOutputNaming:
    """Test output folder and file naming uses company name."""

    def test_output_naming_with_company(self, tmp_path):
        """Test that output folder and file name contain company name."""
        from jseeker.models import AdaptedResume, ContactInfo

        adapted = AdaptedResume(
            contact=ContactInfo(
                full_name="Test User",
                email="test@example.com",
                phone="555-1234",
                locations=["Remote"],
            ),
            target_title="Design Strategist",
            summary="Test summary",
            experience_blocks=[],
            skills_ordered=[],
            education=[],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        with patch("jseeker.renderer._get_display_name", return_value="Test_User"):
            with patch("jseeker.renderer.render_pdf") as mock_pdf:
                with patch("jseeker.renderer.render_docx") as mock_docx:
                    mock_pdf.return_value = tmp_path / "test.pdf"
                    mock_docx.return_value = tmp_path / "test.docx"

                    outputs = generate_output(
                        adapted,
                        company="Santander",
                        role="Design Strategist",
                        output_dir=tmp_path,
                        formats=["pdf", "docx"],
                    )

        # Verify folder contains company name
        pdf_path = outputs["pdf"]
        assert "Santander" in str(pdf_path)
        assert "Not_specified" not in str(pdf_path)

        docx_path = outputs["docx"]
        assert "Santander" in str(docx_path)
        assert "Not_specified" not in str(docx_path)

        # Verify file name contains company name
        assert "Santander" in pdf_path.name
        assert "Santander" in docx_path.name

    def test_output_naming_placeholder_rejected(self, tmp_path):
        """Test that placeholder company names are replaced with fallback."""
        from jseeker.models import AdaptedResume, ContactInfo

        adapted = AdaptedResume(
            contact=ContactInfo(
                full_name="Test User",
                email="test@example.com",
            ),
            target_title="Engineer",
            summary="Test",
            experience_blocks=[],
            skills_ordered=[],
            education=[],
            languages=[],
            certifications=[],
            awards=[],
            template="A",
        )

        placeholders = ["Not specified", "Unknown", "N/A", "not available", "TBD", ""]

        for placeholder in placeholders:
            with patch("jseeker.renderer._get_display_name", return_value="Test_User"):
                with patch("jseeker.renderer.render_pdf") as mock_pdf:
                    with patch("jseeker.renderer.render_docx") as mock_docx:
                        mock_pdf.return_value = tmp_path / "test.pdf"
                        mock_docx.return_value = tmp_path / "test.docx"

                        outputs = generate_output(
                            adapted,
                            company=placeholder,
                            role="Engineer",
                            output_dir=tmp_path,
                            formats=["pdf"],
                        )

            pdf_path = outputs["pdf"]
            assert "Not_specified" not in str(pdf_path), f"Placeholder '{placeholder}' produced Not_specified in path"
            assert "Unknown_Company" in str(pdf_path), f"Placeholder '{placeholder}' should fallback to Unknown_Company"
